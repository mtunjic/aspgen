from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "aspgen-architecture-blueprint.docx"
ASSETS = ROOT / "architecture-assets"
ASSETS.mkdir(exist_ok=True)

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
GOLD = "B7831E"
GRAY = "5B6573"
INK = "182230"


def font_file():
    for candidate in (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/calibri.ttf")):
        if candidate.exists():
            return str(candidate)
    return None


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", style=None, before=0, after=6, line=1.10, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    run = p.add_run(text)
    set_run(run, size=9.2, color="273444", font="Consolas")
    return p


def callout(doc, label, text, fill=PALE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    fixed_table(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_run(r, size=10.5, color=label_color, bold=True)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)
    paragraph(doc, after=4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    fixed_table(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=9.5, color="FFFFFF", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if len(table.rows) % 2 == 0:
                shade(cells[i], "FFFFFF")
            else:
                shade(cells[i], PALE)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=9.4, color=INK)
    paragraph(doc, after=8)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def diagram_gateway(path):
    image = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(image)
    font_path = font_file()
    font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(font_path, 23) if font_path else ImageFont.load_default()
    boxes = [(45, 300, 390, 500, "CLI GATEWAY\naspgen new / add", NAVY), (500, 160, 875, 310, "Profile selector\napp • backend • UI", BLUE), (500, 390, 875, 540, "Manifest boundary\ncomponents + state", BLUE), (990, 100, 1330, 245, "Simple\nActive Record", GOLD), (990, 310, 1330, 455, "DDD / CQRS\nClean layers", BLUE), (990, 520, 1330, 665, "Prism WPF\nMVVM modules", "557A95"), (1450, 300, 1750, 500, "Generated\nsolution +\nprojects", NAVY)]
    for x1, y1, x2, y2, label, color in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#" + color, outline="#" + color, width=3)
        draw.multiline_text(((x1 + x2) / 2, (y1 + y2) / 2), label, font=font, fill="white", anchor="mm", align="center", spacing=6)
    arrows = [((390, 400), (500, 235)), ((390, 400), (500, 465)), ((875, 235), (990, 175)), ((875, 465), (990, 385)), ((875, 465), (990, 595)), ((1330, 175), (1450, 390)), ((1330, 385), (1450, 400)), ((1330, 595), (1450, 420))]
    for (x1, y1), (x2, y2) in arrows:
        draw.line((x1, y1, x2, y2), fill="#7890A8", width=6)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill="#7890A8")
    draw.text((900, 760), "One command surface • multiple architecture outcomes • incremental regeneration", font=small, fill="#4D5A68", anchor="mm")
    image.save(path)


def diagram_flow(path):
    image = Image.new("RGB", (1800, 620), "white")
    draw = ImageDraw.Draw(image)
    font_path = font_file()
    font = ImageFont.truetype(font_path, 26) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(font_path, 23) if font_path else ImageFont.load_default()
    labels = [(45, 250, 320, 390, "HTTP / WPF\nrequest", NAVY), (400, 250, 700, 390, "Endpoint /\nViewModel", BLUE), (780, 250, 1050, 390, "Handler /\nStore", BLUE), (1130, 250, 1400, 390, "Domain\nrule", GOLD), (1480, 250, 1755, 390, "EF Core /\nSQLite default", NAVY)]
    for x1, y1, x2, y2, label, color in labels:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#" + color, outline="#" + color, width=3)
        draw.multiline_text(((x1 + x2) / 2, (y1 + y2) / 2), label, font=font, fill="white", anchor="mm", align="center", spacing=5)
    for i in range(len(labels) - 1):
        x1, y1, x2, y2, *_ = labels[i]
        nx1, ny1, *_ = labels[i + 1]
        draw.line((x2, (y1 + y2) / 2, nx1, (ny1 + y2) / 2), fill="#7890A8", width=6)
        draw.polygon([(nx1, (ny1 + y2) / 2), (nx1 - 18, (ny1 + y2) / 2 - 10), (nx1 - 18, (ny1 + y2) / 2 + 10)], fill="#7890A8")
    draw.text((900, 90), "Vertical slice keeps change local; dependency direction stays stable", font=small, fill="#" + NAVY, anchor="mm")
    draw.text((900, 520), "Result<T> • HTTP status • observable boundary", font=small, fill="#4D5A68", anchor="mm")
    image.save(path)


def build():
    gateway = ASSETS / "gateway.png"
    flow = ASSETS / "flow.png"
    diagram_gateway(gateway)
    diagram_flow(flow)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, "1F4D78", 8, 4)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Running furniture.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("aspgen  |  Architecture Blueprint"), size=8.5, color=GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Architecture Theta  •  Version 1.0"), size=8.5, color=GRAY)

    # Cover / masthead.
    paragraph(doc, "ARCHITECTURE BLUEPRINT", after=12)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.runs[0], size=11, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    set_run(title.add_run("aspgen"), size=34, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    set_run(subtitle.add_run("Gatewayed generation for scalable .NET systems"), size=17, color=GRAY)
    meta = add_table(doc, ["Document", "Audience", "Status"], [["Architecture Blueprint", "Engineering teams and maintainers", "First release baseline"]], [2600, 4200, 2560])
    for cell in meta.rows[0].cells:
        shade(cell, NAVY)
    callout(doc, "THESIS", "aspgen is a controlled gateway: one command surface composes a project profile, records intent in a manifest, and emits only the layers that the selected architecture can own.", fill="E8EEF5", label_color=NAVY)
    paragraph(doc, "This blueprint explains the theory behind the generator, the boundaries it protects, and three end-to-end paths from command to running application.", after=12)
    paragraph(doc, "Prepared 03 August 2026  •  Based on the Architecture Theta ruleset", after=18)
    doc.add_page_break()

    heading(doc, "1. Executive intent", 1)
    paragraph(doc, "The generator is designed for teams that want speed without surrendering architectural ownership. It treats scaffolding as a bounded design activity: a profile selects the architectural grammar, a manifest records the decision, and incremental commands extend the same grammar without regenerating the whole solution.")
    callout(doc, "DECISION", "The generator should optimize for safe composition, not maximum file output. A generated file that has no owning project is a defect, even if its template rendered successfully.")
    add_bullet(doc, "Clean Architecture protects dependency direction: Domain <- Application <- Infrastructure <- WebApi.")
    add_bullet(doc, "Vertical slices keep a feature's handler, validator, endpoint, and DTOs close together.")
    add_bullet(doc, "DDD makes contexts and aggregate roots explicit where business invariants justify the cost.")
    add_bullet(doc, "Prism modules keep desktop UI composition local, discoverable, and testable.")
    add_bullet(doc, "The simple profile is intentionally smaller: direct EF Core CRUD for Rails-like throughput.")

    heading(doc, "2. The gatewayed architecture", 1)
    paragraph(doc, "A gatewayed generator is neither a passive template copier nor a full application framework. It is a policy boundary. Every command passes through four decisions: target application, backend profile, UI profile, and incremental state.")
    doc.add_picture(str(gateway), width=Inches(6.5))
    cap = paragraph(doc, "Figure 1. The command gateway turns profile intent into an owned solution graph.", after=8)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run(run, size=9, color=GRAY, italic=True)
    add_table(doc, ["Gateway decision", "Question", "Result"], [
        ("Application", "What host is being built?", "webapi, wpf, blazor, or fullstack"),
        ("Backend", "How much architecture is justified?", "simple Active Record or DDD/CQRS"),
        ("UI", "Which composition model owns screens?", "Prism/DryIoc modules and MVVM"),
        ("State", "What already exists?", ".aspgen/manifest.json and markers"),
    ], [1800, 3500, 4060])

    heading(doc, "3. Patterns and theory", 1)
    paragraph(doc, "The generator combines patterns that solve different failure modes. The important design move is not to treat them as interchangeable: each pattern has a territory, and the gateway decides when that territory exists.")
    add_table(doc, ["Pattern / theory", "Problem addressed", "Generator expression"], [
        ("Clean Architecture", "Framework and persistence details leaking into policy", "Project references point inward; Domain has no outward dependency."),
        ("DDD bounded context", "A large model becoming an undifferentiated vocabulary", "Contexts scope aggregates, repositories, services, and events."),
        ("Aggregate root", "Invalid state being changed from arbitrary callers", "Commands enter through the root and domain guards protect invariants."),
        ("CQRS + vertical slice", "Features becoming cross-layer scavenger hunts", "Each slice owns command/query, handler, validator, and endpoint."),
        ("Prism MVVM", "Desktop views knowing application orchestration", "Views bind to ViewModels; modules register navigation and regions."),
        ("Active Record", "Overengineering a small CRUD surface", "Simple models and direct EF endpoints reduce ceremony by design."),
        ("Progressive delivery", "Regeneration destroying local work", "Manifest-aware add commands mutate only the requested capability."),
    ], [2100, 3150, 4110])
    callout(doc, "RULE OF THUMB", "Use DDD when invariants and language boundaries are the product. Use the simple profile when CRUD throughput and low ceremony are the product. Do not use complexity as a substitute for a domain decision.", fill="FFF8E8", label_color=GOLD)

    heading(doc, "4. Blueprint of a generated solution", 1)
    add_table(doc, ["Capability", "Owned by", "Typical generated artifacts"], [
        ("Domain", "Domain project", "Entities, aggregate behavior, guards, value objects"),
        ("Application", "Application project", "Commands, queries, handlers, validators, Result<T>"),
        ("Infrastructure", "Infrastructure project", "EF Core context, repositories, audit interceptor, SQLite by default / PostgreSQL opt-in"),
        ("HTTP host", "WebApi project", "Minimal API endpoints, OpenAPI, Scalar, health checks"),
        ("Desktop UI", "Desktop project", "Prism modules, views, ViewModels, typed events, API stores"),
        ("Manifest", ".aspgen directory", "Project identity, components, contexts, backend and seed profile"),
    ], [1800, 2200, 5360])
    paragraph(doc, "The solution file lists projects; SDK-style project files automatically include source and XAML below those project roots. The generator therefore treats project ownership—not manual file enumeration—as the visibility invariant.")

    heading(doc, "5. Three examples", 1)
    heading(doc, "Example A — Simple fullstack with dummy data", 2)
    paragraph(doc, "Use this path for a classic CRUD product, internal tool, or prototype where the shortest safe path to a working screen matters more than domain ceremony.")
    code_block(doc, "aspgen new Library --app fullstack --simple --seed:dummy --theme:wpfui\naspgen add entity Book title:string author:string pages:int published:date available:bool --project Library")
    paragraph(doc, "The gateway emits one Web API project and one Desktop project. The entity becomes an EF Core model, direct CRUD endpoints, a WPF module, and three deterministic sample records created at API startup.")
    add_table(doc, ["Layer", "Output"], [
        ("WebApi", "Models/Book.cs, Data/AppDbContext.cs, Features/Book/BookEndpoints.cs"),
        ("Desktop", "Modules/Book with list/edit form and API-backed store"),
        ("Seed", "Data/DatabaseSeeder.cs with typed sample values"),
    ], [1800, 7560])

    heading(doc, "Example B — DDD/CQRS fullstack", 2)
    paragraph(doc, "Use this path when the entity carries business meaning, invariants, and a lifecycle that should remain explicit. The backend is split into Domain, Application, Infrastructure, and WebApi projects.")
    code_block(doc, "aspgen new Commerce --app fullstack --backend:ddd --seed dummy\naspgen add entity Order customerName:string total:decimal submitted:datetime paid:bool --project Commerce")
    paragraph(doc, "The entity command produces a repository contract, EF repository, DbSet registration, separate CQRS commands and queries, FluentValidation, Minimal API endpoints, a WPF module, and startup seed data. Handler discovery registers IHandler<TRequest,TResponse> implementations automatically.")
    callout(doc, "DEPENDENCY TEST", "A generated DDD project is healthy only when Domain compiles without Application or Infrastructure references, and the WebApi host composes the layers at the edge.")

    heading(doc, "Example C — Add UI to an existing Web API", 2)
    paragraph(doc, "Use incremental generation when the server exists first or when a team wants to introduce desktop workflows without regenerating backend work.")
    code_block(doc, "aspgen new Operations --app webapi --backend ddd\naspgen add entity Person name:string active:bool --project Operations\naspgen add ui --theme:wpfui --project Operations\naspgen add module Reports --project Operations")
    paragraph(doc, "The UI command renders the Desktop project, rewrites the solution membership, and preserves the existing Web API project graph. The module command then adds a Prism module and registers it in App.xaml.cs. No unrelated layers are regenerated.")

    heading(doc, "6. Runtime flow", 1)
    paragraph(doc, "The same architectural idea appears at runtime: the request crosses a narrow gateway, enters a local slice, applies a domain rule, and reaches persistence only through an owned abstraction.")
    doc.add_picture(str(flow), width=Inches(6.5))
    cap = paragraph(doc, "Figure 2. Runtime flow from UI or HTTP boundary to persistence and back.", after=8)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run(run, size=9, color=GRAY, italic=True)
    add_number(doc, "Map transport input into a request record or ViewModel form model.")
    add_number(doc, "Validate at the application boundary; keep endpoint code thin.")
    add_number(doc, "Execute a handler or API store operation with async database/HTTP calls.")
    add_number(doc, "Enforce invariants in the aggregate or domain entity when DDD is selected.")
    add_number(doc, "Return Result<T>, a typed response, or an appropriate HTTP status code.")

    heading(doc, "7. Incremental safety model", 1)
    paragraph(doc, "Incremental generation is a controlled mutation protocol. The manifest answers what the project is; marker comments answer where a new registration belongs; filesystem tests answer whether the output has an owner.")
    add_table(doc, ["Safety mechanism", "Why it exists", "Failure prevented"], [
        ("Manifest components", "Persist selected profiles and generated capabilities", "A DDD command being applied to a simple project"),
        ("Host markers", "Insert registrations deterministically", "Duplicate or misplaced endpoint/module registrations"),
        ("Refuse overwrite", "Protect user-owned files by default", "Silent loss of local implementation"),
        ("Project ownership tests", "Verify paths against actual solution roots", "Orphan Domain files in WPF-only projects"),
        ("Generated build checks", "Compile representative outputs", "Stale DI registrations and template drift"),
    ], [2200, 3400, 3760])
    callout(doc, "QUALITY GATE", "A scaffold is not complete when the command exits successfully. It is complete when the generated files have an owning project, the solution includes that project, and the representative output builds.")

    heading(doc, "8. Recommended operating model", 1)
    add_number(doc, "Choose the smallest profile that can express the current use case.")
    add_number(doc, "Generate the host and backend once; add entities and modules incrementally.")
    add_number(doc, "Run migrations and database updates as an explicit deployment step.")
    add_number(doc, "Treat --seed:dummy as development-only data; do not enable it for production environments.")
    add_number(doc, "Keep custom changes in source files and use exported templates for organization-wide defaults.")
    add_number(doc, "Run go test ./..., go vet ./..., template validation, and at least one generated dotnet build before release.")

    heading(doc, "9. Closing theory", 1)
    paragraph(doc, "The architecture is intentionally asymmetric. Domain code is conservative and inward-facing; WebApi and WPF are composition edges; the generator itself is a policy gateway that keeps those edges aligned. This makes the system scalable in two directions: simple projects can stay simple, while domain-heavy projects can grow without abandoning the same command vocabulary.")
    callout(doc, "NORTH STAR", "Generate less, own more: every emitted file should have a reason, a project, a boundary, and a path back to a tested runtime behavior.", fill="E8EEF5", label_color=NAVY)

    heading(doc, "Appendix — profile quick reference", 1)
    add_table(doc, ["Command", "Best fit", "Primary result"], [
        ("new NAME --app webapi", "Clean server baseline", "Domain/Application/Infrastructure/WebApi"),
        ("new NAME --simple", "Classic CRUD", "Single EF Core WebApi project"),
        ("new NAME --backend ddd", "Business-rich domain", "DDD + CQRS + SQLite API; PostgreSQL opt-in"),
        ("new NAME --app wpf", "Desktop-first", "Prism/DryIoc WPF shell"),
        ("add ui", "Web API first, desktop later", "Desktop project added to solution"),
        ("add entity", "Vertical slice increment", "Backend CRUD, WPF module, or both"),
        ("--seed dummy", "Development demo data", "Typed deterministic startup records"),
    ], [2600, 2760, 4000])
    paragraph(doc, "Reference basis: Architecture Theta in doc/architecture-theta.md; Prism WPF conventions; Clean Architecture, DDD, CQRS, Vertical Slice, MVVM, and Active Record patterns.", after=0)

    doc.core_properties.title = "aspgen Architecture Blueprint"
    doc.core_properties.subject = "Gatewayed architecture for generated .NET and Prism applications"
    doc.core_properties.author = "aspgen Architecture Theta"
    doc.core_properties.keywords = "aspgen, Clean Architecture, DDD, CQRS, Prism, WPF, scaffolding"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
